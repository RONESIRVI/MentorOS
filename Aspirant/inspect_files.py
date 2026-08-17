import pandas as pd
from docx import Document

with open('output.txt', 'w', encoding='utf-8') as f:
    f.write("--- EXCEL FILE ---\n")
    try:
        df = pd.read_excel(r"D:\Journals\RSCDC  Recovery\RSCDC Recovery Report.xlsx")
        f.write(f"Columns: {df.columns.tolist()}\n")
        f.write(df.head(5).to_string() + "\n")
    except Exception as e:
        f.write(f"Error reading Excel: {e}\n")

    f.write("\n--- WORD DOCUMENT ---\n")
    try:
        doc = Document(r"D:\Journals\RSCDC  Recovery\RSCDC Recovery Report.docx")
        f.write("Paragraphs:\n")
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip():
                f.write(f"[{i}]: {p.text}\n")

        f.write("\nTables:\n")
        for i, table in enumerate(doc.tables):
            f.write(f"Table {i}:\n")
            for r_idx, row in enumerate(table.rows):
                row_data = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                f.write(f"  Row {r_idx}: {row_data}\n")
    except Exception as e:
        f.write(f"Error reading Word: {e}\n")
